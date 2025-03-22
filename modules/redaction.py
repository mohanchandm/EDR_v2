from docx import Document
import streamlit as st
import pdfplumber
from io import BytesIO
import docx2txt
import re
from mistralai import Mistral
import base64
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import io
from typing import List, Dict, Optional, Tuple

# Mistral API key (load from Streamlit secrets)
MISTRAL_API_KEY = st.secrets.get("mistral_api_key", "")  # Default to hardcoded if not in secrets
mistral_client = Mistral(api_key=MISTRAL_API_KEY)

def redact_text(text: str, entities: List[Dict[str, any]]) -> str:
    """
    Redact text based on detected entities, handling overlapping spans.

    Args:
        text (str): The input text to redact.
        entities (List[Dict[str, any]]): List of entities with 'start', 'end', 'label', etc.

    Returns:
        str: The redacted text.
    """
    # Sort entities by start position and length (longer first to handle overlaps)
    entities = sorted(entities, key=lambda x: (x["start"], -(x["end"] - x["start"])))
    redacted_text = text
    offset = 0  # Track offset due to redactions

    for entity in entities:
        start, end = entity["start"] + offset, entity["end"] + offset
        entity_text = redacted_text[start:end]
        redacted_label = f"[REDACTED {entity['label'].upper()}]"
        # Calculate padding to preserve length
        padding = " " * max(0, len(entity_text) - len(redacted_label))
        # Replace the entity text
        redacted_text = redacted_text[:start] + redacted_label + padding + redacted_text[end:]
        # Update offset based on length difference
        offset += len(redacted_label + padding) - len(entity_text)

    return redacted_text

def clean_ocr_text(text: str) -> str:
    """
    Clean OCR artifacts from text.

    Args:
        text (str): The input text to clean.

    Returns:
        str: The cleaned text.
    """
    # Fix repetitive character removal (e.g., "aaa" -> "a")
    text = re.sub(r'([a-zA-Z])\1{2,}', r'\1', text)  # Only letters, 2+ repetitions
    # Normalize spaces
    text = re.sub(r'\s+', ' ', text)
    # Fix date placeholders
    text = re.sub(r'\$(\d{4}-\d{2}-\d{2})\$', r'\1', text)
    # Collapse multiple newlines
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def encode_image_to_base64(file_obj: BytesIO) -> str:
    """
    Encode an uploaded image file to base64.

    Args:
        file_obj (BytesIO): The uploaded image file.

    Returns:
        str: The base64-encoded string.
    """
    image_data = file_obj.read()
    return base64.b64encode(image_data).decode('utf-8')

def extract_text_from_image(file_obj: BytesIO) -> str:
    """
    Extract markdown text from an image using Mistral OCR.

    Args:
        file_obj (BytesIO): The uploaded image file.

    Returns:
        str: The extracted markdown text.
    """
    base64_image = encode_image_to_base64(file_obj)
    mime_type = "image/jpeg" if file_obj.name.endswith('.jpg') else "image/png"
    image_url = f"data:{mime_type};base64,{base64_image}"
    
    ocr_response = mistral_client.ocr.process(
        model="mistral-ocr-latest",
        document={
            "type": "image_url",
            "image_url": image_url
        }
    )
    
    markdown_text = ""
    for page in ocr_response.pages:
        markdown_text += page.markdown + "\n\n"
    return clean_ocr_text(markdown_text)

def extract_text_from_file(file_obj: BytesIO) -> Optional[str]:
    """
    Extract text from a file, using OCR for images and low-confidence PDFs.

    Args:
        file_obj (BytesIO): The uploaded file.

    Returns:
        Optional[str]: The extracted text, or None if extraction fails.
    """
    try:
        if file_obj.name.endswith('.txt'):
            text = file_obj.read().decode('utf-8', errors='replace')
            return clean_ocr_text(text)
        
        elif file_obj.name.endswith('.docx'):
            text = docx2txt.process(file_obj)
            text = re.sub(r'\n{2,}', '\n', text)
            return clean_ocr_text(text)
        
        elif file_obj.name.endswith('.pdf'):
            text = ''
            with pdfplumber.open(file_obj) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text(layout=True)
                    # Check if text extraction is low-confidence (e.g., empty or too short)
                    if not page_text or len(page_text.strip()) < 10:
                        # Fallback to OCR for this page
                        file_obj.seek(0)  # Reset file pointer
                        page_image = page.to_image(resolution=300).original
                        buffer = BytesIO()
                        page_image.save(buffer, format="PNG")
                        buffer.name = "temp.png"
                        buffer.seek(0)
                        page_text = extract_text_from_image(buffer)
                    if page_text:
                        page_text = clean_ocr_text(page_text)
                        text += f"\n[Page {page_num + 1}]\n{page_text}"
                    else:
                        text += f"\n[Page {page_num + 1}] - Unable to extract text"
            return text.strip()
        
        elif file_obj.name.endswith(('.jpg', '.png')):
            return extract_text_from_image(file_obj)
        
        else:
            return None
    
    except Exception as e:
        raise Exception(f"Error extracting text from {file_obj.name}: {str(e)}")

def markdown_to_table_data(markdown_text: str) -> List[List[str]]:
    """
    Parse markdown table into a list of lists for rendering, with optimized detection.

    Args:
        markdown_text (str): The markdown text containing a table.

    Returns:
        List[List[str]]: The parsed table data as a list of rows.
    """
    lines = markdown_text.split('\n')
    table_data = []
    
    for line in lines:
        if line.strip().startswith('|') and line.strip().endswith('|') and not line.strip().startswith('| ---'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
        elif table_data:  # If we've already started parsing a table, stop at the first non-table line
            break
    
    return table_data

def markdown_to_docx(markdown_text: str) -> BytesIO:
    """
    Convert markdown text with tables to a docx file.

    Args:
        markdown_text (str): The markdown text to convert.

    Returns:
        BytesIO: The generated docx file as a BytesIO object.
    """
    doc = DocxDocument()
    table_data = markdown_to_table_data(markdown_text)
    
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell
    
    remaining_text = '\n'.join(line for line in markdown_text.split('\n') if not line.startswith('|'))
    for line in remaining_text.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def markdown_to_pdf(markdown_text: str) -> BytesIO:
    """
    Convert markdown text with tables to a PDF file.

    Args:
        markdown_text (str): The markdown text to convert.

    Returns:
        BytesIO: The generated PDF file as a BytesIO object.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    table_data = markdown_to_table_data(markdown_text)
    if table_data:
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(table)
    
    remaining_text = '\n'.join(line for line in markdown_text.split('\n') if not line.startswith('|'))
    for line in remaining_text.split('\n'):
        if line.strip():
            story.append(Paragraph(line.strip(), styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def markdown_to_image(markdown_text: str) -> BytesIO:
    """
    Convert markdown text with tables to an image.

    Args:
        markdown_text (str): The markdown text to convert.

    Returns:
        BytesIO: The generated image file as a BytesIO object.
    """
    table_data = markdown_to_table_data(markdown_text)
    num_rows = len(table_data) if table_data else 0
    num_cols = len(table_data[0]) if table_data and table_data[0] else 0
    cell_width = 150
    cell_height = 30
    width = max(num_cols * cell_width, 600)
    height = max(num_rows * cell_height + (len(markdown_text.split('\n')) - num_rows) * 20 + 50, 400)
    
    img = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 14)
    except:
        font = ImageFont.load_default()
    
    y = 10
    if table_data:
        for i, row in enumerate(table_data):
            x = 10
            for j, cell in enumerate(row):
                draw.rectangle([x, y, x + cell_width, y + cell_height], outline='black')
                draw.text((x + 5, y + 5), cell, font=font, fill='black')
                x += cell_width
            y += cell_height
    
    remaining_text = '\n'.join(line for line in markdown_text.split('\n') if not line.startswith('|'))
    for line in remaining_text.split('\n'):
        if line.strip():
            draw.text((10, y), line.strip(), font=font, fill='black')
            y += 20
    
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer

def convert_to_original_format(redacted_text: str, original_filename: str) -> BytesIO:
    """
    Convert redacted markdown to the original file format.

    Args:
        redacted_text (str): The redacted text to convert.
        original_filename (str): The original filename to determine the format.

    Returns:
        BytesIO: The converted file as a BytesIO object.
    """
    ext = original_filename.split('.')[-1].lower()
    if ext == 'txt':
        return BytesIO(redacted_text.encode('utf-8'))
    elif ext == 'docx':
        return markdown_to_docx(redacted_text)
    elif ext == 'pdf':
        return markdown_to_pdf(redacted_text)
    elif ext in ('jpg', 'png'):
        return markdown_to_image(redacted_text)
    return BytesIO(redacted_text.encode('utf-8'))

def redact_file_content(file_obj: BytesIO, model: any, labels: List[str], threshold: float) -> Tuple[Optional[str], Optional[str], List[Dict[str, any]]]:
    """
    Redact file content and return both original and redacted versions.

    Args:
        file_obj (BytesIO): The uploaded file.
        model (any): The model for entity prediction.
        labels (List[str]): The labels to detect.
        threshold (float): The confidence threshold for entity detection.

    Returns:
        Tuple[Optional[str], Optional[str], List[Dict[str, any]]]: The original text, redacted text, and detected entities.
    """
    original_text = extract_text_from_file(file_obj)
    if original_text:
        entities = model.predict_entities(original_text, labels, threshold=threshold)
        redacted_text = redact_text(original_text, entities)
        return original_text, redacted_text, entities
    return None, None, []
